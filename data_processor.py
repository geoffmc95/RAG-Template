"""
Dynamic Data Processing System for Generic RAG Template
Handles various file types and automatically processes them for RAG ingestion.
"""

import os
import pandas as pd
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
import mimetypes
from llama_index.core import Document
from llama_index.core.readers import SimpleDirectoryReader
from llama_parse import LlamaParse
import json
import docx
import csv
from config import RAGConfig, DataType, get_supported_file_extensions

class DataProcessor:
    """Dynamic data processor that handles multiple file types."""
    
    def __init__(self, config: RAGConfig, llamacloud_api_key: str = None):
        self.config = config
        self.llamacloud_api_key = llamacloud_api_key or os.getenv("LLAMACLOUD_API_KEY")
        self.supported_extensions = get_supported_file_extensions()
        
    def detect_file_type(self, file_path: Path) -> Optional[DataType]:
        """Detect the data type of a file based on its extension."""
        file_extension = file_path.suffix.lower()
        
        for data_type, extensions in self.supported_extensions.items():
            if file_extension in extensions:
                return DataType(data_type)
        
        return None
    
    def get_files_in_directory(self, directory: str = None) -> Dict[DataType, List[Path]]:
        """Get all supported files in the data directory, grouped by type."""
        data_dir = Path(directory or self.config.data_directory)
        if not data_dir.exists():
            return {}
        
        files_by_type = {}
        
        for file_path in data_dir.rglob("*"):
            if file_path.is_file():
                file_type = self.detect_file_type(file_path)
                if file_type and file_type.value in self.config.supported_file_types:
                    if file_type not in files_by_type:
                        files_by_type[file_type] = []
                    files_by_type[file_type].append(file_path)
        
        return files_by_type
    
    def process_pdf_files(self, file_paths: List[Path]) -> List[Document]:
        """Process PDF files using LlamaParse."""
        if not self.llamacloud_api_key:
            print("Warning: No LlamaCloud API key found. Using basic PDF processing.")
            return self._process_pdfs_basic(file_paths)
        
        try:
            parser = LlamaParse(
                api_key=self.llamacloud_api_key,
                result_type="markdown",
                verbose=True,
                language="en",
                parsing_instruction=self.config.domain_config.parsing_instructions
            )
            
            documents = []
            for file_path in file_paths:
                print(f"Processing PDF: {file_path.name}")
                file_documents = parser.load_data(str(file_path))
                for doc in file_documents:
                    doc.metadata.update({
                        "file_name": file_path.name,
                        "file_type": "pdf",
                        "file_path": str(file_path)
                    })
                documents.extend(file_documents)
            
            return documents
        except Exception as e:
            print(f"Error with LlamaParse: {e}. Falling back to basic PDF processing.")
            return self._process_pdfs_basic(file_paths)
    
    def _process_pdfs_basic(self, file_paths: List[Path]) -> List[Document]:
        """Basic PDF processing fallback."""
        documents = []
        for file_path in file_paths:
            try:
                reader = SimpleDirectoryReader(input_files=[str(file_path)])
                file_documents = reader.load_data()
                for doc in file_documents:
                    doc.metadata.update({
                        "file_name": file_path.name,
                        "file_type": "pdf",
                        "file_path": str(file_path)
                    })
                documents.extend(file_documents)
            except Exception as e:
                print(f"Error processing PDF {file_path.name}: {e}")
        return documents
    
    def process_excel_files(self, file_paths: List[Path]) -> List[Document]:
        """Process Excel files."""
        documents = []
        
        for file_path in file_paths:
            try:
                print(f"Processing Excel: {file_path.name}")
                
                # Read all sheets
                excel_file = pd.ExcelFile(file_path)
                
                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    # Convert DataFrame to text representation
                    content = f"Sheet: {sheet_name}\n\n"
                    content += f"Summary: {len(df)} rows, {len(df.columns)} columns\n\n"
                    content += f"Columns: {', '.join(df.columns.tolist())}\n\n"
                    
                    # Add data preview
                    content += "Data Preview:\n"
                    content += df.head(10).to_string(index=False)
                    
                    # Add summary statistics for numeric columns
                    numeric_cols = df.select_dtypes(include=['number']).columns
                    if len(numeric_cols) > 0:
                        content += "\n\nSummary Statistics:\n"
                        content += df[numeric_cols].describe().to_string()
                    
                    # Create document
                    doc = Document(
                        text=content,
                        metadata={
                            "file_name": file_path.name,
                            "file_type": "excel",
                            "file_path": str(file_path),
                            "sheet_name": sheet_name,
                            "rows": len(df),
                            "columns": len(df.columns)
                        }
                    )
                    documents.append(doc)
                    
            except Exception as e:
                print(f"Error processing Excel file {file_path.name}: {e}")
        
        return documents
    
    def process_csv_files(self, file_paths: List[Path]) -> List[Document]:
        """Process CSV files."""
        documents = []
        
        for file_path in file_paths:
            try:
                print(f"Processing CSV: {file_path.name}")
                
                df = pd.read_csv(file_path)
                
                # Convert DataFrame to text representation
                content = f"CSV File: {file_path.name}\n\n"
                content += f"Summary: {len(df)} rows, {len(df.columns)} columns\n\n"
                content += f"Columns: {', '.join(df.columns.tolist())}\n\n"
                
                # Add data preview
                content += "Data Preview:\n"
                content += df.head(10).to_string(index=False)
                
                # Add summary statistics for numeric columns
                numeric_cols = df.select_dtypes(include=['number']).columns
                if len(numeric_cols) > 0:
                    content += "\n\nSummary Statistics:\n"
                    content += df[numeric_cols].describe().to_string()
                
                # Create document
                doc = Document(
                    text=content,
                    metadata={
                        "file_name": file_path.name,
                        "file_type": "csv",
                        "file_path": str(file_path),
                        "rows": len(df),
                        "columns": len(df.columns)
                    }
                )
                documents.append(doc)
                
            except Exception as e:
                print(f"Error processing CSV file {file_path.name}: {e}")
        
        return documents
    
    def process_text_files(self, file_paths: List[Path]) -> List[Document]:
        """Process text files."""
        documents = []
        
        for file_path in file_paths:
            try:
                print(f"Processing text file: {file_path.name}")
                
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                doc = Document(
                    text=content,
                    metadata={
                        "file_name": file_path.name,
                        "file_type": "txt",
                        "file_path": str(file_path),
                        "char_count": len(content)
                    }
                )
                documents.append(doc)
                
            except Exception as e:
                print(f"Error processing text file {file_path.name}: {e}")
        
        return documents
    
    def process_docx_files(self, file_paths: List[Path]) -> List[Document]:
        """Process DOCX files."""
        documents = []
        
        for file_path in file_paths:
            try:
                print(f"Processing DOCX: {file_path.name}")
                
                doc_file = docx.Document(file_path)
                content = ""
                
                for paragraph in doc_file.paragraphs:
                    content += paragraph.text + "\n"
                
                # Also extract text from tables
                for table in doc_file.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            row_text.append(cell.text.strip())
                        content += " | ".join(row_text) + "\n"
                
                doc = Document(
                    text=content,
                    metadata={
                        "file_name": file_path.name,
                        "file_type": "docx",
                        "file_path": str(file_path),
                        "char_count": len(content)
                    }
                )
                documents.append(doc)
                
            except Exception as e:
                print(f"Error processing DOCX file {file_path.name}: {e}")
        
        return documents
    
    def process_json_files(self, file_paths: List[Path]) -> List[Document]:
        """Process JSON files."""
        documents = []
        
        for file_path in file_paths:
            try:
                print(f"Processing JSON: {file_path.name}")
                
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                # Convert JSON to readable text
                content = f"JSON File: {file_path.name}\n\n"
                content += json.dumps(data, indent=2, ensure_ascii=False)
                
                doc = Document(
                    text=content,
                    metadata={
                        "file_name": file_path.name,
                        "file_type": "json",
                        "file_path": str(file_path),
                        "char_count": len(content)
                    }
                )
                documents.append(doc)
                
            except Exception as e:
                print(f"Error processing JSON file {file_path.name}: {e}")
        
        return documents
    
    def process_all_files(self, directory: str = None) -> List[Document]:
        """Process all supported files in the data directory."""
        files_by_type = self.get_files_in_directory(directory)
        all_documents = []
        
        print(f"Found files: {sum(len(files) for files in files_by_type.values())} total")
        
        # Process each file type
        for file_type, file_paths in files_by_type.items():
            print(f"\nProcessing {len(file_paths)} {file_type.value} files...")
            
            if file_type == DataType.PDF:
                documents = self.process_pdf_files(file_paths)
            elif file_type == DataType.EXCEL:
                documents = self.process_excel_files(file_paths)
            elif file_type == DataType.CSV:
                documents = self.process_csv_files(file_paths)
            elif file_type == DataType.TXT:
                documents = self.process_text_files(file_paths)
            elif file_type == DataType.DOCX:
                documents = self.process_docx_files(file_paths)
            elif file_type == DataType.JSON:
                documents = self.process_json_files(file_paths)
            elif file_type == DataType.MARKDOWN:
                documents = self.process_text_files(file_paths)  # Treat markdown as text
            else:
                print(f"Unsupported file type: {file_type}")
                continue
            
            all_documents.extend(documents)
            print(f"Processed {len(documents)} documents from {file_type.value} files")
        
        print(f"\nTotal documents created: {len(all_documents)}")
        return all_documents
    
    def get_processing_summary(self, directory: str = None) -> Dict[str, Any]:
        """Get a summary of files that would be processed."""
        files_by_type = self.get_files_in_directory(directory)
        
        summary = {
            "total_files": sum(len(files) for files in files_by_type.values()),
            "files_by_type": {},
            "supported_types": self.config.supported_file_types,
            "data_directory": directory or self.config.data_directory
        }
        
        for file_type, file_paths in files_by_type.items():
            summary["files_by_type"][file_type.value] = {
                "count": len(file_paths),
                "files": [fp.name for fp in file_paths]
            }
        
        return summary
